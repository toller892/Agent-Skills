#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档解析器模块
支持 txt、md、docx 格式的文档解析
"""

import os
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, field
import logging

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class DocumentSection:
    """文档段落类"""

    title: str
    content: str
    level: int = 1
    keywords: List[str] = field(default_factory=list)


@dataclass
class ParsedDocument:
    """解析后的文档类"""

    title: str
    sections: List[DocumentSection]
    full_text: str
    keywords: List[str] = field(default_factory=list)
    summary: str = ""


class DocumentParser:
    """文档解析器主类"""

    def __init__(self):
        self.supported_formats = [".txt", ".md", ".docx"]

    def parse(self, file_path: str) -> ParsedDocument:
        """
        解析文档

        Args:
            file_path: 文档路径

        Returns:
            ParsedDocument: 解析后的文档对象
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        if path.suffix not in self.supported_formats:
            raise ValueError(
                f"不支持的文件格式: {path.suffix}. 支持格式: {self.supported_formats}"
            )

        logger.info(f"开始解析文档: {file_path}")

        if path.suffix == ".txt":
            return self._parse_txt(path)
        elif path.suffix == ".md":
            return self._parse_md(path)
        elif path.suffix == ".docx":
            return self._parse_docx(path)

        # 确保所有代码路径都有返回值
        raise ValueError(f"不支持的文件格式: {path.suffix}")

    def _parse_txt(self, path: Path) -> ParsedDocument:
        """解析TXT文档"""
        logger.info("解析TXT格式文档")

        with open(path, "r", encoding="utf-8") as f:
            content = f.read()

        # 提取标题（第一行作为标题）
        lines = content.split("\n")
        title = lines[0].strip() if lines else path.stem

        # 分割段落
        sections = self._split_into_sections(content)

        # 提取关键词
        keywords = self._extract_keywords(content)

        return ParsedDocument(
            title=title,
            sections=sections,
            full_text=content,
            keywords=keywords,
            summary=self._generate_summary(content),
        )

    def _parse_md(self, path: Path) -> ParsedDocument:
        """解析Markdown文档"""
        logger.info("解析Markdown格式文档")

        with open(path, "r", encoding="utf-8") as f:
            content = f.read()

        # 提取标题
        title_match = re.search(r"^#\s+(.+)", content, re.MULTILINE)
        title = title_match.group(1).strip() if title_match else path.stem

        # 解析Markdown结构
        sections = self._parse_md_sections(content)

        # 提取关键词
        keywords = self._extract_keywords(content)

        return ParsedDocument(
            title=title,
            sections=sections,
            full_text=content,
            keywords=keywords,
            summary=self._generate_summary(content),
        )

    def _parse_docx(self, path: Path) -> ParsedDocument:
        """解析DOCX文档"""
        logger.info("解析DOCX格式文档")

        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.warning("python-docx 未安装，尝试使用基础解析")
            return self._basic_docx_parse(path)

        doc = DocxDocument(path)

        # 提取标题
        title = ""
        full_text = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                full_text.append(text)
                if not title and text:
                    title = text

        if not title:
            title = path.stem

        content = "\n".join(full_text)

        # 分割段落
        sections = self._split_into_sections(content)

        # 提取关键词
        keywords = self._extract_keywords(content)

        return ParsedDocument(
            title=title,
            sections=sections,
            full_text=content,
            keywords=keywords,
            summary=self._generate_summary(content),
        )

    def _basic_docx_parse(self, path: Path) -> ParsedDocument:
        """基础DOCX解析（使用docling库）"""
        try:
            from docling import Document

            doc = Document(str(path))
            content = doc.get_text()

            title = path.stem
            sections = self._split_into_sections(content)
            keywords = self._extract_keywords(content)

            return ParsedDocument(
                title=title,
                sections=sections,
                full_text=content,
                keywords=keywords,
                summary=self._generate_summary(content),
            )
        except ImportError:
            raise ImportError("请安装 docling: pip install docling")

    def _parse_md_sections(self, content: str) -> List[DocumentSection]:
        """解析Markdown文档的章节结构"""
        sections = []
        lines = content.split("\n")

        current_section = None
        current_content = []

        for line in lines:
            # 检测标题行
            header_match = re.match(r"^(#{1,6})\s+(.+)", line)
            if header_match:
                # 保存之前的段落
                if current_section:
                    current_section.content = "\n".join(current_content)
                    sections.append(current_section)

                level = len(header_match.group(1))
                title = header_match.group(2).strip()
                current_section = DocumentSection(title=title, content="", level=level)
                current_content = []
            elif current_section:
                current_content.append(line)
            else:
                # 标题之前的内容作为引言
                if line.strip():
                    if not sections:
                        sections.append(
                            DocumentSection(title="引言", content=line, level=1)
                        )

        # 保存最后一个段落
        if current_section:
            current_section.content = "\n".join(current_content)
            sections.append(current_section)

        # 如果没有检测到章节，整个内容作为一个章节
        if not sections:
            sections.append(DocumentSection(title="内容", content=content, level=1))

        return sections

    def _split_into_sections(
        self, content: str, num_sections: int = 4
    ) -> List[DocumentSection]:
        """将内容分割成指定数量的章节"""
        # 尝试按段落分割
        paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]

        if len(paragraphs) <= num_sections:
            return [
                DocumentSection(title=f"第{i + 1}部分", content=para, level=1)
                for i, para in enumerate(paragraphs)
            ]

        # 计算每部分应该包含的段落数
        section_size = len(paragraphs) // num_sections
        sections = []

        for i in range(num_sections):
            start_idx = i * section_size
            end_idx = (
                (i + 1) * section_size if i < num_sections - 1 else len(paragraphs)
            )

            section_content = "\n\n".join(paragraphs[start_idx:end_idx])
            section_title = self._generate_section_title(
                paragraphs[start_idx : start_idx + 3]
            )

            sections.append(
                DocumentSection(title=section_title, content=section_content, level=1)
            )

        return sections

    def _generate_section_title(self, content_snippet: List[str]) -> str:
        """根据内容片段生成章节标题"""
        # 提取第一句或第一行的关键词
        for text in content_snippet:
            text = text.strip()
            if len(text) > 10:
                # 提取第一句话
                sentence_end = text.find(".")
                if sentence_end == -1:
                    sentence_end = min(50, len(text))
                return text[:sentence_end].strip()
        return f"部分内容"

    def _extract_keywords(self, content: str) -> List[str]:
        """提取文档关键词"""
        # 移除Markdown格式
        clean_text = re.sub(r"[#*`\[\]()\]]", "", content)

        # 提取常见关键词模式
        words = re.findall(r"\b[a-zA-Z\u4e00-\u9fff]{2,15}\b", clean_text)

        # 统计词频
        word_count = {}
        for word in words:
            word_lower = word.lower()
            word_count[word_lower] = word_count.get(word_lower, 0) + 1

        # 获取出现频率最高的词
        sorted_words = sorted(word_count.items(), key=lambda x: x[1], reverse=True)

        # 过滤常用词
        stop_words = {
            "the",
            "a",
            "an",
            "is",
            "are",
            "was",
            "were",
            "be",
            "been",
            "being",
            "have",
            "has",
            "had",
            "do",
            "does",
            "did",
            "will",
            "would",
            "could",
            "should",
            "may",
            "might",
            "must",
            "shall",
            "这",
            "那",
            "是",
            "在",
            "有",
            "和",
            "与",
            "为",
            "以",
            "于",
        }

        keywords = []
        for word, count in sorted_words:
            if word not in stop_words and count >= 2:
                keywords.append(word)
                if len(keywords) >= 10:
                    break

        return keywords

    def _generate_summary(self, content: str, max_length: int = 200) -> str:
        """生成文档摘要"""
        # 移除Markdown格式
        clean_text = re.sub(r"[#*`\[\]()\]]", "", content)

        # 取前几句话作为摘要
        sentences = re.split(r"[.!?。！？]", clean_text)
        summary_parts = []

        for sentence in sentences:
            sentence = sentence.strip()
            if sentence and len(sentence) > 20:
                summary_parts.append(sentence)
                if sum(len(p) for p in summary_parts) > max_length:
                    break

        summary = " ".join(summary_parts)

        if len(summary) > max_length:
            summary = summary[:max_length] + "..."

        return summary

    def split_for_twitter(
        self, document: ParsedDocument, num_cards: int = 4
    ) -> List[Dict]:
        """
        将文档分割成适合Twitter分享的卡片内容

        Args:
            document: 解析后的文档
            num_cards: 卡片数量，默认4张

        Returns:
            List[Dict]: 卡片内容列表
        """
        cards = []

        # 获取文档标题和摘要
        title = document.title
        summary = document.summary

        # 分割主要内容
        sections = document.sections
        total_sections = len(sections)
        sections_per_card = (total_sections + num_cards - 1) // num_cards

        for i in range(num_cards):
            start_idx = i * sections_per_card
            end_idx = min((i + 1) * sections_per_card, total_sections)

            card_sections = sections[start_idx:end_idx]

            # 构建卡片内容 - 即使sections为空也创建卡片
            card_content = {
                "card_number": i + 1,
                "total_cards": num_cards,
                "title": title,
                "subtitle": f"第 {i + 1} 部分 / 共 {num_cards} 部分"
                if num_cards > 1
                else None,
                "sections": [
                    {
                        "title": sec.title,
                        "content": self._truncate_content(sec.content, 300),
                    }
                    for sec in card_sections
                ] if card_sections else [],
                "keywords": document.keywords,
                "summary": summary if i == 0 else None,  # 只在第一张卡片显示摘要
                "image_prompt": self._generate_image_prompt(
                    card_sections if card_sections else [], document.keywords
                ) if card_sections else f"Illustration for part {i+1} about {title}",
            }

            cards.append(card_content)

        return cards

    def _truncate_content(self, content: str, max_length: int) -> str:
        """截断内容到指定长度"""
        if len(content) <= max_length:
            return content

        # 在句子边界截断
        truncated = content[:max_length]
        last_period = truncated.rfind(".")
        last_newline = truncated.rfind("\n")

        cut_point = max(last_period, last_newline)
        if cut_point > max_length * 0.5:
            truncated = truncated[: cut_point + 1]
        else:
            truncated = truncated.rstrip() + "..."

        return truncated

    def _generate_image_prompt(
        self, sections: List[DocumentSection], keywords: List[str]
    ) -> str:
        """生成图片生成提示词"""
        # 提取章节标题作为图片主题
        titles = [sec.title for sec in sections[:2]]

        # 结合关键词
        combined_keywords = " ".join(keywords[:5])

        prompt = f"""
        Create a visually stunning, modern illustration representing: {", ".join(titles)}
        
        Style requirements:
        - Professional and modern design
        - Vivid colors with gradient effects
        - Clean and minimal aesthetic
        - Suitable for social media sharing
        - High resolution (16:9 aspect ratio)
        
        Key visual elements: {combined_keywords}
        """

        return prompt.strip()


# 测试代码
if __name__ == "__main__":
    parser = DocumentParser()

    # 测试解析
    try:
        # 创建一个测试文档
        test_content = """
# 测试文档标题

这是文档的第一部分内容。我们正在测试文档解析器的功能。

## 第一章 简介

本章介绍了文档的基本结构。文档解析器需要能够识别不同的格式和内容层次。

### 1.1 文档格式

支持TXT、Markdown和DOCX格式。每种格式都有其特点。

### 1.2 解析目标

我们的目标是提取文档的关键信息，并生成适合社交媒体分享的内容。

## 第二章 详细分析

这一章节包含了更详细的内容分析。

### 2.1 技术实现

使用Python实现文档解析功能。

### 2.2 应用场景

文档解析可以应用于多种场景。
        """

        # 保存测试文件
        test_file = Path("test_document.md")
        with open(test_file, "w", encoding="utf-8") as f:
            f.write(test_content)

        # 解析文档
        doc = parser.parse("test_document.md")

        print(f"文档标题: {doc.title}")
        print(f"章节数量: {len(doc.sections)}")
        print(f"关键词: {doc.keywords}")
        print(f"摘要: {doc.summary}")

        # 测试Twitter分割
        cards = parser.split_for_twitter(doc)
        print(f"\nTwitter卡片数量: {len(cards)}")
        for i, card in enumerate(cards):
            print(f"\n卡片 {i + 1}:")
            print(f"  标题: {card['title']}")
            print(f"  小标题: {card['subtitle']}")
            print(f"  章节数: {len(card['sections'])}")

        # 清理测试文件
        test_file.unlink()

    except Exception as e:
        logger.error(f"测试失败: {e}")
